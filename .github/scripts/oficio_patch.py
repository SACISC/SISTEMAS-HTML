"""
Motor reutilizable para llenar oficios de Word a partir de una plantilla,
preservando 100% el diseno (membrete, fuentes, tamanos).

Dos operaciones:
  1. reemplazar_textos(): sustituye texto exacto en document.xml (para
     campos de un solo valor: fecha, folio, destinatario, asunto, etc.)
  2. reconstruir_tabla(): reemplaza las filas de datos de una tabla por
     una cantidad variable de filas nuevas, clonando el formato de la
     fila de referencia (para listas de folios de longitud variable).
"""
import zipfile
import shutil
import copy
import docx


def reemplazar_textos(docx_in_path, docx_out_path, reemplazos: dict):
    """Copia el docx y sustituye texto exacto dentro de word/document.xml.
    reemplazos = { "texto original exacto": "texto nuevo" }
    Lanza ValueError si algun texto original no se encuentra (para nunca
    fallar en silencio y dejar un oficio a medio llenar)."""
    with zipfile.ZipFile(docx_in_path) as z:
        xml = z.read('word/document.xml').decode('utf-8')

    faltantes = []
    for viejo, nuevo in reemplazos.items():
        if viejo not in xml:
            faltantes.append(viejo)
            continue
        xml = xml.replace(viejo, nuevo)

    if faltantes:
        raise ValueError("No se encontraron estos textos en la plantilla: " + " | ".join(faltantes[:5]))

    shutil.copy(docx_in_path, docx_out_path)
    with zipfile.ZipFile(docx_in_path) as zin:
        with zipfile.ZipFile(docx_out_path, 'w', zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename == 'word/document.xml':
                    data = xml.encode('utf-8')
                zout.writestr(item, data)


def reconstruir_tabla(docx_path, tabla_idx, fila_referencia_idx, filas_nuevas):
    """Abre un docx YA con los textos sustituidos, y reemplaza las filas de
    datos de la tabla [tabla_idx] (a partir de fila_referencia_idx) por
    'filas_nuevas' (lista de listas de texto, una lista por fila).
    Clona el formato de la fila de referencia para cada fila nueva, asi que
    la tabla crece o se achica segun cuantos folios traiga la semana."""
    d = docx.Document(docx_path)
    tabla = d.tables[tabla_idx]
    fila_ref = tabla.rows[fila_referencia_idx]

    # Quitar todas las filas de datos existentes (conservando encabezado(s))
    for row in tabla.rows[fila_referencia_idx:]:
        row._element.getparent().remove(row._element)

    for valores in filas_nuevas:
        nueva = copy.deepcopy(fila_ref._element)
        tabla._element.append(nueva)
        fila_obj = tabla.rows[-1]
        for celda, valor in zip(fila_obj.cells, valores):
            for p in celda.paragraphs:
                if p.runs:
                    p.runs[0].text = str(valor)
                    for r in p.runs[1:]:
                        r.text = ''
                else:
                    p.add_run(str(valor))

    d.save(docx_path)
